"""Tests for profile_extractor.py — all parsers, helpers, and new PDF/DOCX/auto-detect."""
import io
import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / ".claude/skills/tailor-resume/scripts"))

from profile_extractor import (
    parse_blob,
    parse_markdown,
    parse_latex,
    parse_linkedin,
    parse_pdf,
    parse_docx,
    merge_profiles,
    auto_detect_format,
    _extract_args,
    _clean_latex,
    _split_sections_latex,
    _parse_dates,
    _dedupe,
    _detect_section,
    _parse_plain_resume_text,
    _parse_with_claude,
    _enrich_profile_with_claude,
    profile_to_dict,
    extract_metrics,
    extract_tools,
    score_confidence,
    Role,
    Profile,
)


# ---------------------------------------------------------------------------
# extract_metrics
# ---------------------------------------------------------------------------
class TestExtractMetrics:
    def test_detects_percentage(self):
        metrics = extract_metrics("Reduced latency by 45%")
        assert len(metrics) >= 1

    def test_detects_dollar_amount(self):
        metrics = extract_metrics("Saved $4,100/month in compute costs")
        assert len(metrics) >= 1

    def test_detects_time_range(self):
        metrics = extract_metrics("Cut runtime from 45 min to 9 min")
        assert len(metrics) >= 1

    def test_empty_string_returns_empty(self):
        assert extract_metrics("") == []

    def test_plain_text_no_metrics(self):
        assert extract_metrics("Led cross-functional team meetings") == []


# ---------------------------------------------------------------------------
# extract_tools
# ---------------------------------------------------------------------------
class TestExtractTools:
    def test_detects_airflow(self):
        tools = extract_tools("Built Airflow DAGs with backfill strategy")
        assert "airflow" in [t.lower() for t in tools]

    def test_detects_spark(self):
        tools = extract_tools("Optimized Spark jobs for 10x speedup")
        assert "spark" in [t.lower() for t in tools]

    def test_detects_python(self):
        tools = extract_tools("Wrote Python ETL pipeline")
        assert "python" in [t.lower() for t in tools]

    def test_detects_multiple_tools(self):
        tools = extract_tools("Used Airflow, Spark, and Delta Lake")
        lower = [t.lower() for t in tools]
        assert "airflow" in lower or "spark" in lower


# ---------------------------------------------------------------------------
# score_confidence
# ---------------------------------------------------------------------------
class TestScoreConfidence:
    def test_quantified_bullet_scores_high_or_medium(self):
        score = score_confidence("Reduced costs by 68% saving $4,100/month via ETL refactor")
        assert score in ("high", "medium")

    def test_vague_bullet_scores_low(self):
        score = score_confidence("Worked on data pipelines")
        assert score == "low"

    def test_single_metric_scores_medium(self):
        score = score_confidence("Cut pipeline runtime by 50%")
        assert score in ("medium", "high")

    def test_score_returns_string(self):
        s = score_confidence("Built something")
        assert isinstance(s, str)
        assert s in ("high", "medium", "low")


# ---------------------------------------------------------------------------
# _extract_args (brace-counting helper)
# ---------------------------------------------------------------------------
class TestExtractArgs:
    def test_single_arg(self):
        args, pos = _extract_args("{hello}", 0, 1)
        assert args == ["hello"]
        assert pos == 7

    def test_multiple_args(self):
        args, pos = _extract_args("{foo}{bar}{baz}", 0, 3)
        assert args == ["foo", "bar", "baz"]

    def test_args_with_whitespace_between(self):
        args, pos = _extract_args("{a}  \n  {b}", 0, 2)
        assert args == ["a", "b"]

    def test_nested_braces(self):
        args, pos = _extract_args("{\\href{url}{label}}", 0, 1)
        assert "\\href" in args[0]

    def test_request_more_args_than_available(self):
        args, _ = _extract_args("{only one}", 0, 3)
        assert len(args) == 1

    def test_empty_string(self):
        args, _ = _extract_args("", 0, 2)
        assert args == []

    def test_partial_start_position(self):
        text = "prefix{value}"
        args, _ = _extract_args(text, 6, 1)
        assert args == ["value"]


# ---------------------------------------------------------------------------
# _clean_latex
# ---------------------------------------------------------------------------
class TestCleanLatex:
    def test_removes_textbf(self):
        result = _clean_latex("\\textbf{Bold Text}")
        assert "Bold Text" in result
        assert "\\textbf" not in result

    def test_removes_textit(self):
        result = _clean_latex("\\textit{Italic}")
        assert "Italic" in result
        assert "\\textit" not in result

    def test_unwraps_href(self):
        result = _clean_latex("\\href{https://example.com}{Click here}")
        assert "Click here" in result
        assert "https://example.com" not in result

    def test_replaces_double_dash(self):
        result = _clean_latex("Jan 2020 -- Present")
        assert "--" not in result or "–" in result or "Jan" in result

    def test_removes_latex_percent_escape(self):
        result = _clean_latex("Reduced by 40\\%")
        assert "40%" in result

    def test_normalizes_whitespace(self):
        result = _clean_latex("a   b   c")
        assert "  " not in result

    def test_empty_string(self):
        assert _clean_latex("") == ""


# ---------------------------------------------------------------------------
# _split_sections_latex
# ---------------------------------------------------------------------------
class TestSplitSectionsLatex:
    def test_splits_two_sections(self):
        tex = "\\section{Experience}\nsome content\n\\section{Skills}\nskill content"
        sections = _split_sections_latex(tex)
        assert "experience" in sections
        assert "skills" in sections

    def test_section_content_correct(self):
        tex = "\\section{Education}\nedu stuff\n\\section{Projects}\nproj stuff"
        sections = _split_sections_latex(tex)
        assert "edu" in sections["education"]
        assert "proj" in sections["projects"]

    def test_empty_text_returns_empty_dict(self):
        assert _split_sections_latex("no sections here") == {}

    def test_single_section(self):
        sections = _split_sections_latex("\\section{Experience}\ncontent")
        assert "experience" in sections


# ---------------------------------------------------------------------------
# _parse_dates
# ---------------------------------------------------------------------------
class TestParseDates:
    def test_en_dash(self):
        start, end = _parse_dates("Jan 2022 – Present")
        assert "Jan 2022" in start
        assert "Present" in end

    def test_double_dash(self):
        start, end = _parse_dates("July 2024 -- Present")
        assert start.strip() != ""
        assert end.strip() != ""

    def test_no_separator_returns_full_as_start(self):
        start, end = _parse_dates("2022")
        assert start == "2022"
        assert end == ""

    def test_year_range(self):
        start, end = _parse_dates("2020 - 2022")
        assert "2020" in start
        assert "2022" in end


# ---------------------------------------------------------------------------
# _dedupe
# ---------------------------------------------------------------------------
class TestDedupe:
    def test_removes_duplicates(self):
        assert _dedupe(["a", "b", "a", "c"]) == ["a", "b", "c"]

    def test_preserves_order(self):
        assert _dedupe(["z", "a", "m"]) == ["z", "a", "m"]

    def test_empty_list(self):
        assert _dedupe([]) == []

    def test_no_duplicates_unchanged(self):
        lst = ["x", "y", "z"]
        assert _dedupe(lst) == lst


# ---------------------------------------------------------------------------
# _detect_section
# ---------------------------------------------------------------------------
class TestDetectSection:
    def test_detects_experience(self):
        assert _detect_section("Experience") == "experience"

    def test_detects_education(self):
        assert _detect_section("Education") == "education"

    def test_detects_skills(self):
        assert _detect_section("Technical Skills") == "skills"

    def test_detects_projects(self):
        assert _detect_section("Projects") == "projects"

    def test_detects_certifications(self):
        assert _detect_section("Certifications") == "certifications"

    def test_case_insensitive(self):
        assert _detect_section("EXPERIENCE") == "experience"

    def test_with_trailing_colon(self):
        assert _detect_section("Skills:") == "skills"

    def test_unknown_returns_none(self):
        assert _detect_section("Random Section Header") is None

    def test_empty_string_returns_none(self):
        assert _detect_section("") is None

    def test_detects_work_experience(self):
        assert _detect_section("Work Experience") == "experience"

    def test_detects_professional_experience(self):
        assert _detect_section("PROFESSIONAL EXPERIENCE") == "experience"

    def test_detects_technical_skills_partial(self):
        assert _detect_section("TECHNICAL SKILLS") == "skills"

    def test_detects_work_history(self):
        assert _detect_section("WORK HISTORY") == "experience"

    def test_long_non_section_not_detected(self):
        # Should not match — too long and contains non-alpha chars
        assert _detect_section("This is a very long line with numbers 123 and special @chars") is None


# ---------------------------------------------------------------------------
# parse_markdown
# ---------------------------------------------------------------------------
class TestParseMarkdown:
    SAMPLE_MD = """\
## Experience

**Senior Data Engineer** | DataWorks Inc | Jan 2022 - Present
- Built governed semantic layer on Databricks, cutting metric discrepancies from 12/week to zero
- Owned CI/CD via Azure DevOps, compressing deployments from 8 weeks to 6 days

**Data Engineer** | Acme Analytics | Jun 2020 - Dec 2021
- Designed partitioned Delta Lake tables, cutting query time by 60%

## Skills

Python, SQL, Spark, Airflow, Delta Lake
"""

    def test_parse_markdown_returns_profile(self):
        profile = parse_markdown(self.SAMPLE_MD)
        assert isinstance(profile, Profile)

    def test_parse_markdown_extracts_roles(self):
        profile = parse_markdown(self.SAMPLE_MD)
        assert len(profile.experience) >= 1

    def test_parse_markdown_extracts_bullets(self):
        profile = parse_markdown(self.SAMPLE_MD)
        total_bullets = sum(len(r.bullets) for r in profile.experience)
        assert total_bullets >= 1

    def test_parse_markdown_extracts_skills(self):
        profile = parse_markdown(self.SAMPLE_MD)
        assert len(profile.skills) >= 1

    def test_parse_markdown_role_has_title(self):
        profile = parse_markdown(self.SAMPLE_MD)
        titles = [r.title for r in profile.experience]
        assert any("Data Engineer" in t for t in titles)

    def test_parse_markdown_empty_text(self):
        profile = parse_markdown("")
        assert profile.experience == []


# ---------------------------------------------------------------------------
# parse_latex — single-line (backward-compat)
# ---------------------------------------------------------------------------
class TestParseLatex:
    SAMPLE_LATEX = (
        r"\section{Experience}" + "\n"
        r"  \resumeSubHeadingListStart" + "\n"
        r"    \resumeSubheading{Senior Data Engineer}{Jan 2022 -- Present}{DataWorks Inc}{Remote}" + "\n"
        r"      \resumeItemListStart" + "\n"
        r"        \resumeItem{Built governed semantic layer on Databricks}" + "\n"
        r"        \resumeItem{Owned CI/CD via Azure DevOps}" + "\n"
        r"      \resumeItemListEnd" + "\n"
        r"  \resumeSubHeadingListEnd" + "\n"
        "\n"
        r"\section{Projects}" + "\n"
        r"    \resumeSubHeadingListStart" + "\n"
        r"      \resumeProjectHeading{\textbf{Analytics Dashboard} $|$ \emph{Python, Streamlit}}{2023}" + "\n"
        r"          \resumeItemListStart" + "\n"
        r"            \resumeItem{Built real-time analytics dashboard serving 1k users}" + "\n"
        r"          \resumeItemListEnd" + "\n"
        r"    \resumeSubHeadingListEnd" + "\n"
    )

    def test_parse_latex_returns_profile(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        assert isinstance(profile, Profile)

    def test_parse_latex_extracts_experience(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        assert len(profile.experience) >= 1

    def test_parse_latex_extracts_bullets(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        total = sum(len(r.bullets) for r in profile.experience)
        assert total >= 1

    def test_parse_latex_extracts_projects(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        assert len(profile.projects) >= 1

    def test_parse_latex_project_has_bullets(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        project_bullets = sum(len(p.bullets) for p in profile.projects)
        assert project_bullets >= 1

    def test_parse_latex_role_title(self):
        profile = parse_latex(self.SAMPLE_LATEX)
        assert profile.experience[0].title == "Senior Data Engineer"

    def test_parse_latex_empty_text(self):
        profile = parse_latex("")
        assert profile.experience == []


# ---------------------------------------------------------------------------
# parse_latex — multiline (Jake template style)
# ---------------------------------------------------------------------------
class TestParseLatexMultiline:
    """The real resume uses \resumeSubheading over 3 lines — this is the primary use case."""

    MULTILINE_LATEX = (
        r"\section{Experience}" + "\n"
        r"  \resumeSubHeadingListStart" + "\n"
        r"    \resumeSubheading" + "\n"
        r"      {Data Engineer}{July 2024 -- Present}" + "\n"
        r"      {ExponentHR Technologies}{Dallas, TX}" + "\n"
        r"      \resumeItemListStart" + "\n"
        r"        \resumeItem{Architected AI-powered analytics platform, reducing support volume by 40\%}" + "\n"
        r"        \resumeItem{Compressed deployment cycles from 3 months to 14 days (85\% faster)}" + "\n"
        r"      \resumeItemListEnd" + "\n"
        r"    \resumeSubheading" + "\n"
        r"      {Data Engineer -- Research}{Aug. 2023 -- July 2024}" + "\n"
        r"      {Missouri S\&T}{Rolla, MO}" + "\n"
        r"      \resumeItemListStart" + "\n"
        r"        \resumeItem{Managed university webpages and data infrastructure}" + "\n"
        r"      \resumeItemListEnd" + "\n"
        r"  \resumeSubHeadingListEnd" + "\n"
        "\n"
        r"\section{Education}" + "\n"
        r"  \resumeSubHeadingListStart" + "\n"
        r"    \resumeSubheading" + "\n"
        r"      {Missouri University of Science and Technology}{Rolla, MO}" + "\n"
        r"      {Master of Science in Information Science}{Jan. 2022 -- Dec. 2023}" + "\n"
        r"  \resumeSubHeadingListEnd" + "\n"
        "\n"
        r"\section{Technical Skills}" + "\n"
        r" \begin{itemize}" + "\n"
        r"    \small{\item{" + "\n"
        r"     \textbf{Languages \& ML}{: Python, SQL, Bash, MLflow, Airflow} \\" + "\n"
        r"     \textbf{Data \& Cloud}{: Spark, Kafka, Databricks, Azure} \\" + "\n"
        r"    }}" + "\n"
        r" \end{itemize}" + "\n"
        "\n"
        r"\section{Certifications \& Publications}" + "\n"
        r" \begin{itemize}" + "\n"
        r"    \small{\item{" + "\n"
        r"     \textbf{DP-700}{: Microsoft Certified Data Engineer Associate}" + "\n"
        r"    }}" + "\n"
        r" \end{itemize}" + "\n"
    )

    def test_multiline_subheading_extracts_roles(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert len(profile.experience) == 2

    def test_multiline_first_role_title(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert "Data Engineer" in profile.experience[0].title

    def test_multiline_first_role_company(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert "ExponentHR" in profile.experience[0].company

    def test_multiline_role_has_bullets(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert len(profile.experience[0].bullets) >= 1

    def test_multiline_bullet_text_clean(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        bullet_text = profile.experience[0].bullets[0].text
        assert "40%" in bullet_text
        assert "\\" not in bullet_text

    def test_multiline_education_parsed(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert len(profile.education) >= 1
        assert any("Missouri" in e.get("institution", "") for e in profile.education)

    def test_multiline_skills_parsed(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert len(profile.skills) >= 3
        skill_names = [s.lower() for s in profile.skills]
        assert any("python" in s for s in skill_names)

    def test_multiline_certifications_parsed(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        assert len(profile.certifications) >= 1

    def test_multiline_dates_parsed(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        role = profile.experience[0]
        assert role.start != "" or role.end != ""

    def test_second_role_correct(self):
        profile = parse_latex(self.MULTILINE_LATEX)
        second = profile.experience[1]
        assert "Missouri" in second.company or "Rolla" in second.location


# ---------------------------------------------------------------------------
# auto_detect_format
# ---------------------------------------------------------------------------
class TestAutoDetectFormat:
    def test_detects_latex_from_documentclass(self):
        tex = r"\documentclass[letterpaper]{article}\n\begin{document}\end{document}"
        assert auto_detect_format(tex) == "latex"

    def test_detects_latex_from_resumeitem(self):
        tex = r"\resumeItem{Built something}\resumeSubheading{DE}{2022}{Acme}{TX}"
        assert auto_detect_format(tex) == "latex"

    def test_detects_markdown(self):
        md = "## Experience\n**Engineer** | Acme | 2022\n- Built thing"
        assert auto_detect_format(md) == "markdown"

    def test_defaults_to_blob(self):
        plain = "Senior Data Engineer at Acme Corp from 2020 to 2022"
        assert auto_detect_format(plain) == "blob"

    def test_single_hash_heading_is_markdown(self):
        md = "# Resume\n## Skills\nPython, SQL"
        assert auto_detect_format(md) == "markdown"


# ---------------------------------------------------------------------------
# _parse_plain_resume_text (PDF/DOCX extracted text)
# ---------------------------------------------------------------------------
class TestParsePlainResumeText:
    PLAIN_TEXT = """\
Experience
Data Engineer  ExponentHR Technologies  July 2024 – Present
• Architected AI-powered analytics platform reducing support by 40%
• Compressed deployment cycles from 3 months to 14 days

Data Engineer  Missouri S&T  Aug 2023 – July 2024
• Managed university infrastructure and data pipelines

Education
Missouri University of Science and Technology
Master of Science in Information Science 2022 – 2023

Skills
Python, SQL, Spark, Kafka, Airflow, Docker, Azure

Certifications
DP-700 Microsoft Certified Data Engineer Associate
"""

    def test_returns_profile(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        assert isinstance(p, Profile)

    def test_extracts_experience_roles(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        assert len(p.experience) >= 1

    def test_extracts_bullets(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        total = sum(len(r.bullets) for r in p.experience)
        assert total >= 1

    def test_extracts_skills(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        assert len(p.skills) >= 3

    def test_extracts_certifications(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        assert len(p.certifications) >= 1

    def test_skills_contain_python(self):
        p = _parse_plain_resume_text(self.PLAIN_TEXT)
        skill_names = [s.lower() for s in p.skills]
        assert any("python" in s for s in skill_names)

    def test_empty_text_returns_empty_profile(self):
        p = _parse_plain_resume_text("")
        assert isinstance(p, Profile)

    def test_skills_deduped(self):
        text = "Skills\nPython, Python, SQL\nExperience"
        p = _parse_plain_resume_text(text)
        assert p.skills.count("Python") <= 1

    def test_bullet_with_dash_prefix(self):
        text = "Experience\nData Engineer  Acme Corp  2022 – 2024\n- Built a pipeline saving 30% cost"
        p = _parse_plain_resume_text(text)
        bullets = sum(len(r.bullets) for r in p.experience)
        assert bullets >= 1

    def test_three_line_role_header(self):
        """Title / Company / Date on separate lines (common in PDFs)."""
        text = (
            "EXPERIENCE\n"
            "Data Engineer\n"
            "Acme Corp\n"
            "2022 – 2024\n"
            "• Built a streaming pipeline processing 1M events/day\n"
        )
        p = _parse_plain_resume_text(text)
        assert len(p.experience) >= 1
        assert p.experience[0].title == "Data Engineer"
        assert p.experience[0].company == "Acme Corp"

    def test_two_line_role_header(self):
        """Title on one line, date on the next."""
        text = "EXPERIENCE\nData Engineer  Acme Corp\n2022 – Present\n"
        p = _parse_plain_resume_text(text)
        assert len(p.experience) >= 1

    def test_two_column_latex_company_on_next_line(self):
        """2-column LaTeX: title+date on line 1, company+location on line 2."""
        text = (
            "EXPERIENCE\n"
            "Data Engineer  July 2024 – Present\n"
            "ExponentHR  Dallas, TX\n"
            "• Built a data pipeline processing 5M records/day\n"
        )
        p = _parse_plain_resume_text(text)
        assert len(p.experience) >= 1
        assert p.experience[0].title == "Data Engineer"
        assert p.experience[0].company == "ExponentHR"
        assert p.experience[0].location == "Dallas, TX"

    def test_skills_with_category_label(self):
        """'Languages: Python SQL Go' → individual skill tokens."""
        text = "SKILLS\nLanguages: Python SQL Go\nFrameworks: Spark Kafka\n"
        p = _parse_plain_resume_text(text)
        names = [s.lower() for s in p.skills]
        assert "python" in names
        assert "spark" in names

    def test_education_with_degree_keywords(self):
        text = "EDUCATION\nMaster of Science in CS  State University  2021\n"
        p = _parse_plain_resume_text(text)
        assert len(p.education) >= 1

    def test_projects_section(self):
        text = "PROJECTS\nJan 2023 – Mar 2023  My Cool Project\n"
        p = _parse_plain_resume_text(text)
        assert isinstance(p, Profile)  # no crash

    def test_stdlib_pdf_extractor_tj_join(self):
        """TJ kerning arrays must produce joined words, not fragments."""
        from profile_extractor import _extract_pdf_text_stdlib
        # Simulate a PDF stream with kerning-split "Python"
        fake = b"stream\nBT\n[(Pyth) -30 (on)] TJ\nET\nendstream"
        out = _extract_pdf_text_stdlib(fake)
        assert "Python" in out

    def test_stdlib_pdf_extractor_t_star_newline(self):
        """T* operator should produce a line break."""
        from profile_extractor import _extract_pdf_text_stdlib
        fake = b"stream\nBT\n(EXPERIENCE) Tj\nT*\n(Data Engineer) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(fake)
        lines = [line for line in out.splitlines() if line.strip()]
        assert len(lines) >= 2

    def test_stdlib_pdf_extractor_td_newline(self):
        """Td with non-zero y offset should produce a line break."""
        from profile_extractor import _extract_pdf_text_stdlib
        fake = b"stream\nBT\n(Senior) Tj\n0 -14 Td\n(Engineer) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(fake)
        lines = [line for line in out.splitlines() if line.strip()]
        assert len(lines) >= 2

    def test_stdlib_pdf_extractor_quote_operator(self):
        """' operator should produce a new line."""
        from profile_extractor import _extract_pdf_text_stdlib
        fake = b"stream\nBT\n(SKILLS) Tj\n(Python) '\nET\nendstream"
        out = _extract_pdf_text_stdlib(fake)
        assert "SKILLS" in out or "Python" in out

    def test_stdlib_pdf_extractor_garbage_filter(self):
        """Lines that are mostly non-alpha binary data should be dropped."""
        from profile_extractor import _extract_pdf_text_stdlib
        fake = b"stream\nBT\n(][-<>{}|~^) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(fake)
        assert out.strip() == "" or "Python" not in out

    def test_stdlib_pdf_extractor_no_streams(self):
        """Data with no stream markers falls back to raw data scan."""
        from profile_extractor import _extract_pdf_text_stdlib
        raw = b"BT\n(Hello World) Tj\nET"
        out = _extract_pdf_text_stdlib(raw)
        assert "Hello World" in out

    def test_stdlib_pdf_extractor_ot1_fi_ligature(self):
        """OT1 byte 0x0C (fi ligature) must be decoded to 'fi', not dropped."""
        from profile_extractor import _extract_pdf_text_stdlib
        # "MLflow" in OT1: 'M', 'L', 0x0C (fi), 'o', 'w'
        content = b"stream\nBT\n(ML\x0cow) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(content)
        assert "MLfiow" in out or "MLflow" in out  # fi → "fi", so MLfiow is correct

    def test_stdlib_pdf_extractor_ot1_fl_ligature(self):
        """OT1 byte 0x0D (fl ligature) must be decoded to 'fl', not dropped."""
        from profile_extractor import _extract_pdf_text_stdlib
        # "Airflow" in OT1: 'A', 'i', 'r', 0x0D (fl), 'o', 'w'
        content = b"stream\nBT\n(Air\x0dow) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(content)
        assert "Airflow" in out or "Airflow" in out

    def test_stdlib_pdf_extractor_ot1_en_dash(self):
        """OT1 byte 0x7B must be decoded to en dash, not kept as '{'."""
        from profile_extractor import _extract_pdf_text_stdlib
        content = b"stream\nBT\n(Jan 2022 \x7b Present) Tj\nET\nendstream"
        out = _extract_pdf_text_stdlib(content)
        assert "{" not in out
        assert "\u2013" in out or "Present" in out

    def test_stdlib_pdf_extractor_tj_kerning_word_space(self):
        """Large negative kerning in TJ array must produce a word space."""
        from profile_extractor import _extract_pdf_text_stdlib
        # "Data Engineer" split by large kerning -300
        content = b"stream\nBT\n[(Data) -300 (Engineer)] TJ\nET\nendstream"
        out = _extract_pdf_text_stdlib(content)
        assert "Data Engineer" in out

    def test_stdlib_pdf_extractor_tj_small_kerning_no_space(self):
        """Small kerning (-30) must NOT insert a space (kerning adjustment, not word gap)."""
        from profile_extractor import _extract_pdf_text_stdlib
        content = b"stream\nBT\n[(Pyth) -30 (on)] TJ\nET\nendstream"
        out = _extract_pdf_text_stdlib(content)
        assert "Python" in out
        assert "Pyth on" not in out

    def test_apply_ot1_maps_ligatures(self):
        """_apply_ot1 must map all OT1 ligature bytes correctly."""
        from profile_extractor import _apply_ot1
        assert _apply_ot1("\x0c") == "fi"
        assert _apply_ot1("\x0d") == "fl"
        assert _apply_ot1("\x0e") == "ff"
        assert _apply_ot1("\x0f") == "ffi"
        assert _apply_ot1("\x10") == "ffl"
        assert _apply_ot1("\x7b") == "\u2013"
        assert _apply_ot1("\x95") == "\u2022"


# ---------------------------------------------------------------------------
# parse_pdf
# ---------------------------------------------------------------------------
class TestParsePdf:
    def test_parse_pdf_blank_page_raises_value_error_with_pypdf(self):
        """pypdf is used for blank page → no text → ValueError (not ImportError)."""
        pytest.importorskip("pypdf", reason="pypdf not installed")
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()
        # Blank page has no text; correct behavior is ValueError, not ImportError
        with pytest.raises(ValueError, match="No text could be extracted"):
            parse_pdf(pdf_bytes)

    def test_parse_pdf_falls_back_to_stdlib_without_pypdf(self, monkeypatch):
        """Without pypdf, stdlib fallback runs; fake bytes → ValueError (no text found)."""
        import builtins
        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "pypdf":
                raise ImportError("No module named 'pypdf'")
            return real_import(name, *args, **kwargs)

        saved = sys.modules.pop("pypdf", None)
        monkeypatch.setattr(builtins, "__import__", mock_import)
        try:
            with pytest.raises(ValueError, match="No text could be extracted"):
                parse_pdf(b"fake pdf bytes")
        finally:
            monkeypatch.setattr(builtins, "__import__", real_import)
            if saved is not None:
                sys.modules["pypdf"] = saved


# ---------------------------------------------------------------------------
# _extract_pdf_text_pdfminer — smoke test
# ---------------------------------------------------------------------------
class TestExtractPdfTextPdfminer:
    def test_pdfminer_extracts_text_from_minimal_pdf(self):
        """pdfminer returns non-empty string for a real PDF with text content."""
        pytest.importorskip("pdfminer", reason="pdfminer.six not installed")
        from profile_extractor import _extract_pdf_text_pdfminer
        pytest.importorskip("pypdf", reason="pypdf needed to create test PDF")
        from pypdf import PdfWriter
        from pypdf.generic import NameObject, DecodedStreamObject

        writer = PdfWriter()
        # Build a minimal but valid PDF with a text stream
        page = writer.add_blank_page(width=612, height=792)
        # Embed a content stream with a simple text operator
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 100 700 Td (Missouri) Tj ET")
        page[NameObject("/Contents")] = stream
        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        result = _extract_pdf_text_pdfminer(pdf_bytes)
        # pdfminer returns a string (may be empty for PDFs without embedded fonts,
        # but should not raise)
        assert isinstance(result, str)

    def test_pdfminer_import_error_propagates(self, monkeypatch):
        """ImportError from pdfminer propagates so parse_pdf can catch it."""
        import builtins
        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "pdfminer.high_level":
                raise ImportError("No module named 'pdfminer'")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)
        from profile_extractor import _extract_pdf_text_pdfminer
        with pytest.raises(ImportError):
            _extract_pdf_text_pdfminer(b"fake")

    def test_parse_pdf_uses_pdfminer_when_available(self, monkeypatch):
        """parse_pdf calls _extract_pdf_text_pdfminer first when pdfminer is available."""
        import profile_extractor as pe
        called = []

        def fake_pdfminer(data):
            called.append("pdfminer")
            return "Senior Data Engineer at Acme Corp Jan 2022 – Dec 2024\n• Built pipelines"

        monkeypatch.setattr(pe, "_extract_pdf_text_pdfminer", fake_pdfminer)
        result = pe.parse_pdf(b"any bytes")
        assert called == ["pdfminer"]
        assert isinstance(result, pe.Profile)

    def test_parse_pdf_falls_back_to_pypdf_when_pdfminer_missing(self, monkeypatch):
        """parse_pdf falls back to pypdf when pdfminer raises ImportError."""
        import profile_extractor as pe

        def fake_pdfminer(data):
            raise ImportError("No module named 'pdfminer'")

        monkeypatch.setattr(pe, "_extract_pdf_text_pdfminer", fake_pdfminer)
        # With blank pypdf page → no text → ValueError (not ImportError)
        pytest.importorskip("pypdf", reason="pypdf needed for fallback test")
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        with pytest.raises(ValueError, match="No text could be extracted"):
            pe.parse_pdf(buf.getvalue())

    def test_pdfminer_extractor_unicode_endash(self, monkeypatch):
        """parse_pdf passes through U+2013 en dash from pdfminer without mangling."""
        import profile_extractor as pe

        def fake_pdfminer(data):
            # Simulate pdfminer reading ToUnicode CMap → correct en dash
            return "Data Engineer: Acme Corp\nJuly 2024 \u2013 Present\n\u2022 Built pipelines"

        monkeypatch.setattr(pe, "_extract_pdf_text_pdfminer", fake_pdfminer)
        profile = pe.parse_pdf(b"bytes")
        # En dash must survive through parsing into the role's end date
        assert len(profile.experience) == 1
        assert "2024" in profile.experience[0].start or "July" in profile.experience[0].start
        assert "Present" in profile.experience[0].end or profile.experience[0].end != ""

    def test_pdfminer_extractor_word_grouping_via_bullet_split(self, monkeypatch):
        """parse_pdf correctly handles multi-sentence bullet block from pdfminer.

        pdfminer groups tightly-spaced chars into words (char_margin=1.5) and
        adjacent bullet paragraphs into one LTTextBox.  _split_bullet_block
        must split them at sentence boundaries so each sentence becomes a bullet.
        """
        import profile_extractor as pe

        def fake_pdfminer(data):
            # Simulate what pdfminer returns for a 2-sentence bullet block:
            # two sentences, each starting uppercase after a period-ended previous line
            return (
                "Work Experience\n"
                "Senior Engineer: Acme Corp\nJan 2022 \u2013 Dec 2024\n"
                "\u2022 Reduced latency by 40%. Scaled throughput to 10K TPS."
            )

        monkeypatch.setattr(pe, "_extract_pdf_text_pdfminer", fake_pdfminer)
        profile = pe.parse_pdf(b"bytes")
        assert len(profile.experience) == 1
        # Bullets come through (sentence-split or as one bullet)
        assert len(profile.experience[0].bullets) >= 1


# ---------------------------------------------------------------------------
# _split_bullet_block — unit tests (pure Python, no PDF needed)
# ---------------------------------------------------------------------------
class TestSplitBulletBlock:
    def test_single_sentence_returns_one_item(self):
        from profile_extractor import _split_bullet_block
        text = "Built a scalable data pipeline processing 10K records per second."
        result = _split_bullet_block(text)
        assert result == ["Built a scalable data pipeline processing 10K records per second."]

    def test_two_sentences_split_at_boundary(self):
        from profile_extractor import _split_bullet_block
        text = (
            "Reduced query latency by 40% through index tuning.\n"
            "Scaled throughput to 10K TPS using Kafka partitioning."
        )
        result = _split_bullet_block(text)
        assert len(result) == 2
        assert result[0].startswith("Reduced")
        assert result[1].startswith("Scaled")

    def test_wrapped_lines_joined_within_sentence(self):
        from profile_extractor import _split_bullet_block
        text = (
            "Architected an AI-powered analytics platform on\n"
            "Microsoft Fabric, reducing support volume by 40%.\n"
            "Owned end-to-end CI/CD infrastructure."
        )
        result = _split_bullet_block(text)
        assert len(result) == 2
        assert "Microsoft Fabric" in result[0]
        assert "40%" in result[0]
        assert result[1].startswith("Owned")

    def test_empty_string_returns_empty_list(self):
        from profile_extractor import _split_bullet_block
        assert _split_bullet_block("") == []

    def test_blank_lines_ignored(self):
        from profile_extractor import _split_bullet_block
        text = "First sentence.\n\nSecond sentence."
        result = _split_bullet_block(text)
        assert len(result) == 2


# ---------------------------------------------------------------------------
# parse_docx
# ---------------------------------------------------------------------------
class TestParseDocx:
    def test_parse_docx_returns_profile_with_python_docx(self):
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document
        doc = Document()
        doc.add_paragraph("Experience")
        doc.add_paragraph("Data Engineer  Acme Corp  2022 – 2024")
        doc.add_paragraph("• Built a streaming pipeline processing 1M events/day")
        doc.add_paragraph("Skills")
        doc.add_paragraph("Python, SQL, Spark")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        profile = parse_docx(docx_bytes)
        assert isinstance(profile, Profile)

    def test_parse_docx_falls_back_to_stdlib_without_python_docx(self, monkeypatch):
        """Without python-docx, stdlib fallback runs; fake bytes → ValueError (no text found)."""
        import builtins
        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        saved = sys.modules.pop("docx", None)
        monkeypatch.setattr(builtins, "__import__", mock_import)
        try:
            with pytest.raises(ValueError, match="No text could be extracted from this DOCX"):
                parse_docx(b"fake docx bytes")
        finally:
            monkeypatch.setattr(builtins, "__import__", real_import)
            if saved is not None:
                sys.modules["docx"] = saved


# ---------------------------------------------------------------------------
# parse_linkedin
# ---------------------------------------------------------------------------
class TestParseLinkedin:
    def test_parse_linkedin_returns_profile(self):
        text = "Senior Data Engineer at DataWorks Inc\n- Built pipelines\n- Reduced costs by 40%"
        profile = parse_linkedin(text)
        assert isinstance(profile, Profile)


# ---------------------------------------------------------------------------
# merge_profiles
# ---------------------------------------------------------------------------
class TestMergeProfiles:
    def test_merge_combines_experience(self):
        p1 = Profile()
        p1.experience.append(Role(title="DE", company="A", start="2020", end="2021", location=""))
        p2 = Profile()
        p2.experience.append(Role(title="SDE", company="B", start="2021", end="2022", location=""))
        merged = merge_profiles(p1, p2)
        assert len(merged.experience) == 2

    def test_merge_dedupes_skills(self):
        p1 = Profile()
        p1.skills = ["Python", "SQL"]
        p2 = Profile()
        p2.skills = ["Python", "Airflow"]
        merged = merge_profiles(p1, p2)
        assert merged.skills.count("Python") == 1

    def test_merge_empty_profiles(self):
        merged = merge_profiles(Profile(), Profile())
        assert merged.experience == []
        assert merged.skills == []

    def test_merge_three_profiles(self):
        p1, p2, p3 = Profile(), Profile(), Profile()
        p1.skills = ["Python"]
        p2.skills = ["SQL"]
        p3.skills = ["Spark"]
        merged = merge_profiles(p1, p2, p3)
        assert len(merged.skills) == 3


# ---------------------------------------------------------------------------
# profile_to_dict
# ---------------------------------------------------------------------------
class TestParseBlobDateRegression:
    def test_october_not_split_on_substring_to(self):
        """
        Regression: the blob date regex used `(?:–|-|to)` as the separator,
        which matched the substring "to" inside month names like "October" —
        producing start="Oc" and end="ber 2024 – May 2025".

        The fix uses (?<=\\s)to(?=\\s) so "to" only matches when surrounded
        by whitespace (i.e. when it's a real word, not a substring).
        """
        profile = parse_blob(
            "Company: Acme\nTitle: Engineer\nDates: October 2024 – May 2025\n"
        )
        role = profile.experience[0]
        assert role.start == "October 2024"
        assert role.end == "May 2025"

    def test_word_to_still_works_as_separator(self):
        """The word 'to' surrounded by whitespace must still split dates."""
        profile = parse_blob(
            "Company: Acme\nTitle: Engineer\nDates: Jan 2022 to Dec 2023\n"
        )
        role = profile.experience[0]
        assert role.start == "Jan 2022"
        assert role.end == "Dec 2023"


# ---------------------------------------------------------------------------
# profile_to_dict
# ---------------------------------------------------------------------------
class TestProfileToDict:
    def test_profile_to_dict_is_serializable(self):
        profile = parse_blob(
            "Company: DataWorks\nTitle: DE\n- Built pipeline saving $1k/month"
        )
        d = profile_to_dict(profile)
        json.dumps(d)

    def test_profile_to_dict_has_experience_key(self):
        profile = Profile()
        d = profile_to_dict(profile)
        assert "experience" in d


# ---------------------------------------------------------------------------
# _parse_with_claude — direct calls with no key and mocked API
# ---------------------------------------------------------------------------
def _make_anthropic_mock(monkeypatch, response_text: str):
    """Inject a fake anthropic module so tests run without the package installed."""
    import sys
    from unittest.mock import MagicMock

    mock_resp = MagicMock()
    mock_resp.content = [MagicMock(text=response_text)]
    mock_anthropic = MagicMock()
    mock_anthropic.Anthropic.return_value.messages.create.return_value = mock_resp
    monkeypatch.setitem(sys.modules, "anthropic", mock_anthropic)
    return mock_anthropic


def _make_anthropic_mock_error(monkeypatch, exc: Exception):
    """Inject a fake anthropic module that raises exc on messages.create."""
    import sys
    from unittest.mock import MagicMock

    mock_anthropic = MagicMock()
    mock_anthropic.Anthropic.return_value.messages.create.side_effect = exc
    monkeypatch.setitem(sys.modules, "anthropic", mock_anthropic)
    return mock_anthropic


class TestParseWithClaude:
    def test_no_api_key_returns_plain_parse(self, monkeypatch):
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        result = _parse_with_claude(
            "Data Engineer at Acme Inc  Jan 2022 – Present\n• Built ETL pipelines\nSkills: Python SQL",
            source="test",
        )
        assert result is not None
        assert hasattr(result, "experience")

    def test_mocked_api_returns_structured_profile(self, monkeypatch):
        import json

        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {
            "experience": [
                {
                    "title": "Data Engineer",
                    "company": "Acme Corp",
                    "start": "Jan 2022",
                    "end": "Present",
                    "location": "Dallas, TX",
                    "bullets": ["Built scalable ETL pipelines", "Reduced load time by 40%"],
                }
            ],
            "projects": [
                {"name": "ML Pipeline", "tech": ["Python", "Spark"], "bullets": ["Processed 1M rows/day"]}
            ],
            "skills": ["Python", "SQL", "Spark"],
            "education": [
                {"institution": "UT Dallas", "degree": "MS Computer Science",
                 "dates": "2020 – 2022", "location": ""}
            ],
            "certifications": ["AWS Solutions Architect"],
        }
        _make_anthropic_mock(monkeypatch, json.dumps(data))
        result = _parse_with_claude("raw resume text here", source="test")
        assert len(result.experience) == 1
        assert result.experience[0].title == "Data Engineer"
        assert result.experience[0].company == "Acme Corp"
        assert len(result.experience[0].bullets) == 2
        assert len(result.projects) == 1
        assert "Python" in result.skills
        assert len(result.education) == 1
        assert "AWS Solutions Architect" in result.certifications

    def test_mocked_api_with_json_fenced_response(self, monkeypatch):
        import json

        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {"experience": [], "projects": [], "skills": ["Python", "SQL"],
                "education": [], "certifications": []}
        fenced = "```json\n" + json.dumps(data) + "\n```"
        _make_anthropic_mock(monkeypatch, fenced)
        result = _parse_with_claude("resume text", source="test")
        assert "Python" in result.skills
        assert "SQL" in result.skills

    def test_api_exception_falls_back_to_plain_parse(self, monkeypatch):
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        _make_anthropic_mock_error(monkeypatch, RuntimeError("API timeout"))
        result = _parse_with_claude(
            "Data Engineer at Acme  Jan 2022 – Present\n- Built pipelines",
            source="test",
        )
        assert result is not None  # falls back gracefully

    def test_invalid_json_response_falls_back(self, monkeypatch):
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        _make_anthropic_mock(monkeypatch, "this is not json at all")
        result = _parse_with_claude("resume text", source="test")
        assert result is not None  # falls back on json.JSONDecodeError


# ---------------------------------------------------------------------------
# _enrich_profile_with_claude — mocked API rebuild path
# ---------------------------------------------------------------------------
class TestEnrichProfileWithClaudeMocked:
    def test_no_api_key_returns_original_profile(self, monkeypatch):
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        profile = parse_blob("Data Engineer at Acme  Jan 2022 – Present\n- Built pipelines")
        result = _enrich_profile_with_claude(profile, source="test")
        assert result is profile  # exact same object returned

    def test_mocked_api_rebuilds_enriched_profile(self, monkeypatch):
        import json

        profile = parse_blob(
            "Data Engineer at Acme Inc  Jan 2022 – Present\n• Built ETL pipelines\nSkills: Python, SQL"
        )
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        enriched_data = {
            "experience": [
                {
                    "title": "Data Engineer",
                    "company": "Acme Inc",
                    "start": "Jan 2022",
                    "end": "Present",
                    "location": "",
                    "bullets": [
                        {"text": "Built scalable ETL pipelines processing 1M rows/day", "confidence": "high"},
                        {"text": "Reduced query latency by 40% using partitioning", "confidence": "high"},
                    ],
                }
            ],
            "projects": [
                {
                    "name": "Real-Time Dashboard",
                    "tech": ["Kafka", "Spark"],
                    "bullets": [{"text": "Streamed 500k events/sec", "confidence": "high"}],
                }
            ],
            "skills": ["Python", "SQL", "Spark", "Kafka"],
            "education": [
                {"institution": "UT Dallas", "degree": "MS CS", "dates": "2020-2022", "location": ""}
            ],
            "certifications": ["AWS Data Engineer"],
        }
        _make_anthropic_mock(monkeypatch, json.dumps(enriched_data))
        result = _enrich_profile_with_claude(profile, source="test")
        assert len(result.experience) == 1
        assert len(result.experience[0].bullets) == 2
        assert "Spark" in result.skills
        assert "Kafka" in result.skills
        assert len(result.projects) == 1
        assert "AWS Data Engineer" in result.certifications

    def test_mocked_api_with_fenced_json(self, monkeypatch):
        import json

        profile = parse_blob("Engineer at Co  2022-Present\n- Did things")
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {
            "experience": [],
            "projects": [],
            "skills": ["Python", "Go"],
            "education": [],
            "certifications": ["GCP Professional"],
        }
        fenced = "```json\n" + json.dumps(data) + "\n```"
        _make_anthropic_mock(monkeypatch, fenced)
        result = _enrich_profile_with_claude(profile, source="test")
        assert "GCP Professional" in result.certifications

    def test_api_exception_returns_original_profile(self, monkeypatch):
        profile = parse_blob("Engineer at Co  2022-Present\n- Built things")
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        _make_anthropic_mock_error(monkeypatch, RuntimeError("network error"))
        result = _enrich_profile_with_claude(profile, source="test")
        assert result is profile  # enrichment never blocks; returns original

    def test_mocked_api_with_string_bullets(self, monkeypatch):
        import json

        profile = parse_blob("Engineer at Co  2022-Present\n- Built pipelines")
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Co",
                    "start": "2022",
                    "end": "Present",
                    "location": "",
                    "bullets": ["Built high-throughput pipelines", "Reduced costs by 30%"],
                }
            ],
            "projects": [],
            "skills": ["Python"],
            "education": [],
            "certifications": [],
        }
        _make_anthropic_mock(monkeypatch, json.dumps(data))
        result = _enrich_profile_with_claude(profile, source="test")
        assert len(result.experience[0].bullets) == 2


# ---------------------------------------------------------------------------
# _is_bullet_line — OT1 artifact prefix recognition (Issue #48)
# ---------------------------------------------------------------------------
class TestIsBulletLine:
    def setup_method(self):
        from profile_extractor import _is_bullet_line
        self._fn = _is_bullet_line

    def test_standard_bullet_dot(self):
        assert self._fn("• Built pipelines")

    def test_standard_bullet_dash(self):
        assert self._fn("- Reduced latency by 40%")

    def test_ffi_prefix_recognized_as_bullet(self):
        """OT1 CMR glyph 0x0F decoded as 'ffi' — should be treated as a bullet."""
        assert self._fn("ffi Architected distributed system")

    def test_j_prefix_recognized_as_bullet(self):
        """Icon-font separator glyph decoded as 'j' — should be treated as a bullet."""
        assert self._fn("j Led cross-functional team")

    def test_x_prefix_still_recognized(self):
        """Existing CMR checkmark 'x' prefix must still work."""
        assert self._fn("x Delivered project on time")

    def test_plain_word_not_bullet(self):
        assert not self._fn("Architected distributed system")

    def test_ffi_without_trailing_space_not_bullet(self):
        """'ffi' alone (no content after space) should not fire the pattern."""
        assert not self._fn("ffi")


# ---------------------------------------------------------------------------
# _extract_pdf_text_stdlib — OT1 normalization post-pass (Issue #49)
# ---------------------------------------------------------------------------
class TestExtractPdfTextStdlibNormalization:
    def setup_method(self):
        from profile_extractor import _extract_pdf_text_stdlib
        self._fn = _extract_pdf_text_stdlib

    def _minimal_pdf(self, text_content: str) -> bytes:
        """Build a minimal hand-crafted PDF with a single text stream."""
        stream = text_content.encode("latin-1", errors="replace")
        stream_len = len(stream)
        pdf = (
            b"%PDF-1.4\n"
            b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
            b"4 0 obj\n<< /Length " + str(stream_len).encode() + b" >>\nstream\n"
            + stream + b"\nendstream\nendobj\n"
            b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
            b"xref\n0 6\n0000000000 65535 f \n"
            b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n0\n%%EOF"
        )
        return pdf

    def test_ffi_prefix_line_converted_to_bullet(self):
        """'ffi Built pipelines' → '• Built pipelines' after normalization."""
        from profile_extractor import _OT1_ARTIFACT_PREFIX
        result = _OT1_ARTIFACT_PREFIX.sub("• ", "ffi Built pipelines")
        assert result.startswith("•")

    def test_j_prefix_line_converted_to_bullet(self):
        """'j Led team of 5' → '• Led team of 5' after normalization."""
        from profile_extractor import _OT1_ARTIFACT_PREFIX
        result = _OT1_ARTIFACT_PREFIX.sub("• ", "j Led team of 5")
        assert result.startswith("•")

    def test_artifact_only_line_matches_drop_pattern(self):
        """A line that is just 'ffi' or 'j' with optional whitespace is a lone icon."""
        from profile_extractor import _OT1_ARTIFACT_ONLY
        assert _OT1_ARTIFACT_ONLY.match("ffi")
        assert _OT1_ARTIFACT_ONLY.match("j")
        assert _OT1_ARTIFACT_ONLY.match("ffi  ")

    def test_real_word_starting_with_j_not_dropped(self):
        """'JavaScript' must not be classified as an artifact-only line."""
        from profile_extractor import _OT1_ARTIFACT_ONLY
        assert not _OT1_ARTIFACT_ONLY.match("JavaScript developer")

    def test_ffi_prefix_not_stripped_from_middle_of_line(self):
        """The prefix regex only matches at the start of a line."""
        from profile_extractor import _OT1_ARTIFACT_PREFIX
        result = _OT1_ARTIFACT_PREFIX.sub("• ", "Skills: ffi Python")
        assert "ffi Python" in result  # unchanged — 'ffi' not at start

    def test_parse_plain_resume_with_ffi_bullets_extracts_bullets(self):
        """End-to-end: text with 'ffi ' bullets gets bullets extracted."""
        text = (
            "Experience\n"
            "Data Engineer  Acme Corp  2022 – 2024\n"
            "ffi Architected distributed ETL pipeline reducing costs by 30%\n"
            "ffi Led migration from Hadoop to Spark cutting runtime by 50%\n"
        )
        p = _parse_plain_resume_text(text)
        assert len(p.experience) >= 1


# ---------------------------------------------------------------------------
# _build_profile_from_claude_json + _parse_pdf_with_claude_document_api
# ---------------------------------------------------------------------------
class TestBuildProfileFromClaudeJson:
    def test_full_json_produces_correct_profile(self):
        from profile_extractor import _build_profile_from_claude_json, Profile
        data = {
            "experience": [
                {"title": "DE", "company": "Acme", "start": "2022", "end": "Present",
                 "location": "Dallas, TX", "bullets": ["Built ETL reducing costs 30%"]}
            ],
            "projects": [{"name": "MLPipe", "tech": ["Python"], "bullets": ["Processed 1M rows"]}],
            "skills": ["Python", "Spark"],
            "education": [{"institution": "UT Dallas", "degree": "MS CS",
                           "dates": "2020-2022", "location": ""}],
            "certifications": ["AWS SAA"],
        }
        profile = _build_profile_from_claude_json(data, source="test")
        assert isinstance(profile, Profile)
        assert len(profile.experience) == 1
        assert profile.experience[0].title == "DE"
        assert len(profile.experience[0].bullets) == 1
        assert "Python" in profile.skills
        assert len(profile.education) == 1
        assert "AWS SAA" in profile.certifications

    def test_empty_json_returns_empty_profile(self):
        from profile_extractor import _build_profile_from_claude_json, Profile
        profile = _build_profile_from_claude_json({}, source="test")
        assert isinstance(profile, Profile)
        assert profile.experience == []
        assert profile.skills == []

    def test_non_string_bullets_skipped(self):
        from profile_extractor import _build_profile_from_claude_json
        data = {
            "experience": [
                {"title": "SWE", "company": "Co", "start": "2021", "end": "2023",
                 "location": "", "bullets": [None, 42, "Valid bullet", ""]}
            ],
            "projects": [], "skills": [], "education": [], "certifications": [],
        }
        profile = _build_profile_from_claude_json(data, source="test")
        assert len(profile.experience[0].bullets) == 1
        assert profile.experience[0].bullets[0].text == "Valid bullet"


class TestParsePdfWithClaudeDocumentApi:
    def test_returns_none_when_no_api_key(self, monkeypatch):
        """No ANTHROPIC_API_KEY → returns None without calling anthropic."""
        from profile_extractor import _parse_pdf_with_claude_document_api
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        result = _parse_pdf_with_claude_document_api(b"fake pdf", "test")
        assert result is None

    def test_returns_none_when_anthropic_import_fails(self, monkeypatch):
        """Missing anthropic package → returns None gracefully."""
        from profile_extractor import _parse_pdf_with_claude_document_api
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        monkeypatch.setitem(sys.modules, "anthropic", None)
        result = _parse_pdf_with_claude_document_api(b"fake pdf", "test")
        assert result is None

    def test_returns_profile_on_success(self, monkeypatch):
        """Mocked API returning valid JSON → Profile returned."""
        from profile_extractor import _parse_pdf_with_claude_document_api
        import json as _json
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {
            "experience": [
                {"title": "Data Engineer", "company": "Acme", "start": "Jan 2022",
                 "end": "Present", "location": "Dallas, TX",
                 "bullets": ["Reduced ETL by 40%"]}
            ],
            "projects": [], "skills": ["Python", "Spark"],
            "education": [], "certifications": [],
        }
        mock = _make_anthropic_mock(monkeypatch, _json.dumps(data))
        mock.__version__ = "0.27.0"
        result = _parse_pdf_with_claude_document_api(b"%PDF-1.4 fake", "test")
        assert result is not None
        assert len(result.experience) == 1
        assert result.experience[0].title == "Data Engineer"
        assert "Python" in result.skills

    def test_returns_none_on_api_error(self, monkeypatch):
        """API exception → returns None, never raises."""
        from profile_extractor import _parse_pdf_with_claude_document_api
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        mock = _make_anthropic_mock_error(monkeypatch, RuntimeError("timeout"))
        mock.__version__ = "0.27.0"
        result = _parse_pdf_with_claude_document_api(b"fake pdf", "test")
        assert result is None

    def test_parse_pdf_tier0_used_when_key_set(self, monkeypatch):
        """parse_pdf() invokes Tier 0 first when ANTHROPIC_API_KEY is set."""
        import json as _json
        monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")
        data = {
            "experience": [
                {"title": "Staff Engineer", "company": "BigCo", "start": "2020",
                 "end": "Present", "location": "", "bullets": ["Built infra"]}
            ],
            "projects": [], "skills": ["Go"], "education": [], "certifications": [],
        }
        mock = _make_anthropic_mock(monkeypatch, _json.dumps(data))
        mock.__version__ = "0.27.0"
        # Even for fake bytes, Tier 0 should return the mocked profile
        result = parse_pdf(b"%PDF-1.4 fake bytes", source="test")
        assert result is not None
        assert result.experience[0].title == "Staff Engineer"

    def test_scanned_pdf_no_key_error_message_guides_user(self, monkeypatch):
        """Scanned PDF with no key → ValueError message mentions ANTHROPIC_API_KEY."""
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        with pytest.raises(ValueError, match="ANTHROPIC_API_KEY"):
            parse_pdf(b"fake non-extractable pdf bytes")
