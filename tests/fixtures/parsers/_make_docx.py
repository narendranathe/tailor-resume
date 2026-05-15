"""
_make_docx.py
Generate the canonical sample.docx fixture for parser integration tests.

Run from repo root:
    python tests/fixtures/parsers/_make_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt


HERE = Path(__file__).parent
OUT = HERE / "sample.docx"


def _add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def build():
    doc = Document()

    # Header / contact
    _add_para(doc, "JANE DOE", size=18, bold=True)
    _add_para(doc, "Senior Data Engineer", size=12)
    _add_para(doc, "jane.doe@example.com | 555-123-4567 | Dallas, TX")

    # Experience
    _add_heading(doc, "Professional Experience")

    _add_para(
        doc,
        "Senior Data Engineer, DataWorks Inc                          March 2022 – Present",
        bold=True,
    )
    _add_bullet(
        doc,
        "Built a governed semantic layer on Azure Databricks with Delta Lake, "
        "defining 40+ KPI metrics across Finance, ML, and BI teams; cut metric "
        "discrepancies from 12 weekly incidents to zero",
    )
    _add_bullet(
        doc,
        "Owned CI/CD end-to-end through Azure DevOps for 15 data pipelines, "
        "compressing deployment cycles from 8 weeks to 6 days",
    )
    _add_bullet(
        doc,
        "Reengineered ETL from full-table reloads to CDC incremental merge upserts "
        "on Delta Lake, cutting batch runtime from 45 min to 9 min and compute costs by 68%",
    )
    _add_bullet(
        doc,
        "Migrated Spark jobs to AKS with HPA, consolidating 18 nodes to 4-8 dynamic, "
        "reducing monthly cloud spend by $4,100",
    )

    _add_para(
        doc,
        "Data Engineer, Acme Analytics Corp                       August 2019 – February 2022",
        bold=True,
    )
    _add_bullet(
        doc,
        "Engineered Azure AI Anomaly Detector pipelines achieving 95%+ accuracy "
        "across 200+ monitored metrics",
    )
    _add_bullet(
        doc,
        "Built Elasticsearch enterprise search indexing 150K+ internal documents, "
        "reducing support desk volume by 75%",
    )
    _add_bullet(
        doc,
        "Developed Python ETL framework with Pytest unit tests and GitHub Actions CI, "
        "reducing production data defects by 40% year-over-year",
    )

    # Education
    _add_heading(doc, "Education")
    _add_para(doc, "University of Missouri-Columbia, Columbia, MO")
    _add_para(doc, "Master of Science in Computer Science, August 2017 – May 2019")

    # Skills
    _add_heading(doc, "Skills")
    _add_para(doc, "Languages: Python, SQL, Bash, Java, Scala")
    _add_para(doc, "Data: Spark, Kafka, Airflow, dbt, Delta Lake")
    _add_para(doc, "Cloud: Azure, AWS, Databricks")
    _add_para(doc, "DevOps: Docker, Kubernetes, Terraform, CI/CD")

    # Certifications
    _add_heading(doc, "Certifications")
    _add_bullet(doc, "AWS Certified Solutions Architect — Associate")
    _add_bullet(doc, "Databricks Certified Data Engineer Professional")

    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
