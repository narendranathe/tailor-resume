"""
cover_letter_renderer.py
Generate a 2-paragraph cover letter from a profile + JD gap report.

Structure (max 250 words total):
  Para 1: Hook -- why this specific role + company (drawn from JD signals)
  Para 2: Impact bridge -- 2-3 STAR-compressed achievements mapped to top JD requirements

Export formats:
  .tex   -- always generated (LaTeX source, shares header style with resume)
  .txt   -- plain text (no LaTeX commands; for copy-paste into job portal text boxes)
  .docx  -- via python-docx (optional dep; skipped gracefully if not installed)
  .pdf   -- deferred to Issue #67 (requires pdflatex subprocess)

Generation methods:
  method="claude"    -- LLM writes both paragraphs; falls back to template on exception
  method="template"  -- rule-based, deterministic, zero API keys needed
"""
from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

_SCRIPTS = Path(__file__).parent
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))

from latex_renderer import escape  # noqa: E402

_TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "cover_letter_template.tex"

_MAX_WORDS = 250


@dataclass
class CoverLetterResult:
    tex: str
    txt: str
    docx_path: Optional[str]
    method_used: str
    word_count: int


def build_cover_letter(
    profile_dict: dict,
    report,  # GapReport -- duck-typed to avoid circular import
    header: dict,
    jd_text: str,
    method: str = "claude",
) -> CoverLetterResult:
    """
    Generate a 2-paragraph cover letter.

    Args:
        profile_dict: Profile dict from profile_extractor (experience, projects, skills...).
        report: GapReport with top_missing and recommendations.
        header: Contact info dict (name, email, phone, linkedin, ...).
        jd_text: Full job description text.
        method: "claude" (LLM, falls back to template) or "template" (rule-based).

    Returns:
        CoverLetterResult with .tex, .txt, .docx_path, .method_used, .word_count.
    """
    if method == "claude":
        return _build_claude(profile_dict, report, header, jd_text)
    return _build_template(profile_dict, report, header, jd_text)


def _build_claude(
    profile_dict: dict,
    report,
    header: dict,
    jd_text: str,
) -> CoverLetterResult:
    """Option: Claude writes the paragraphs. Falls back to template on any exception."""
    try:
        import anthropic  # lazy import — keeps module importable without the dep

        top_gaps = []
        if hasattr(report, "top_missing"):
            top_gaps = [g.category for g in report.top_missing[:3]]

        experience_snippet = json.dumps(
            profile_dict.get("experience", [])[:1], indent=2
        )[:600]

        prompt = (
            "Write a 2-paragraph professional cover letter. Max 250 words total. "
            "Do NOT start with 'I am writing to express my interest'. "
            "Be specific, confident, and results-focused.\n\n"
            f"Job Description (excerpt):\n{jd_text[:800]}\n\n"
            f"Candidate Experience:\n{experience_snippet}\n\n"
            f"Top JD gaps to address: {top_gaps}\n\n"
            "Return ONLY the two paragraphs separated by a blank line. "
            "No salutation, no closing, no extra formatting."
        )

        client = anthropic.Anthropic()
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        # Strip any accidental markdown code fences
        raw = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\n?```$", "", raw, flags=re.MULTILINE)
        raw = raw.strip()

        parts = [p.strip() for p in raw.split("\n\n") if p.strip()]
        para1 = parts[0] if len(parts) > 0 else raw
        para2 = parts[1] if len(parts) > 1 else ""

        para1, para2 = _enforce_word_limit(para1, para2, _MAX_WORDS)
        return _assemble(para1, para2, header, method_used="claude")

    except Exception:
        result = _build_template(profile_dict, report, header, jd_text)
        return CoverLetterResult(
            tex=result.tex,
            txt=result.txt,
            docx_path=result.docx_path,
            method_used="template (claude fallback)",
            word_count=result.word_count,
        )


def _build_template(
    profile_dict: dict,
    report,
    header: dict,
    jd_text: str,
) -> CoverLetterResult:
    """Rule-based generation. Deterministic, zero external deps."""
    company = _extract_company_from_jd(jd_text)

    # Para 1: hook from JD signals
    top_keywords: list[str] = []
    if hasattr(report, "top_missing") and report.top_missing:
        top_keywords = list(report.top_missing[0].jd_keywords[:2])

    skill_list = profile_dict.get("skills", [])
    top_skill = skill_list[0] if skill_list else "data engineering"

    if top_keywords:
        para1 = (
            f"I am excited to apply for this role at {company}. "
            f"Your focus on {top_keywords[0]}"
            + (f" and {top_keywords[1]}" if len(top_keywords) > 1 else "")
            + f" aligns directly with my experience building production-grade {top_skill} systems."
        )
    else:
        para1 = (
            f"I am excited to apply for this role at {company}. "
            f"My background in {top_skill} positions me to make an immediate impact "
            f"on your team's data infrastructure and engineering goals."
        )

    # Para 2: 1-2 bullets from first experience role
    experience = profile_dict.get("experience", [])
    bullet_texts: list[str] = []
    if experience:
        bullets = experience[0].get("bullets", [])
        for b in bullets[:2]:
            text = b.get("text", "") if isinstance(b, dict) else str(b)
            if text:
                bullet_texts.append(text)

    role_company = experience[0].get("company", "my previous role") if experience else "my previous role"

    if bullet_texts:
        achievements = " Additionally, ".join(
            f"at {role_company}, I {b.rstrip('.')}." for b in bullet_texts
        )
        para2 = achievements
    else:
        para2 = (
            "Throughout my career I have consistently delivered measurable results "
            "through strong engineering fundamentals, cross-functional collaboration, "
            "and a focus on reliability and performance."
        )

    para1, para2 = _enforce_word_limit(para1, para2, _MAX_WORDS)
    return _assemble(para1, para2, header, method_used="template")


def _extract_company_from_jd(jd_text: str) -> str:
    """Best-effort company name extraction from JD text."""
    patterns = [
        r"(?:at|join|joining)\s+([A-Z][A-Za-z0-9&\s]{2,30}?)(?:\s+as|\s+to|\.|,)",
        r"^([A-Z][A-Za-z0-9&\s]{2,30}?)\s+is\s+(?:hiring|looking|seeking)",
    ]
    for pat in patterns:
        m = re.search(pat, jd_text, re.IGNORECASE | re.MULTILINE)
        if m:
            return m.group(1).strip()
    return "your company"


def _enforce_word_limit(para1: str, para2: str, limit: int) -> tuple[str, str]:
    """Trim combined paragraphs to `limit` words at the nearest sentence boundary."""
    combined = f"{para1}\n\n{para2}"
    words = combined.split()
    if len(words) <= limit:
        return para1, para2

    # Cut at limit words, then walk back to sentence boundary
    chunk = words[:limit]
    text = " ".join(chunk)
    # Find last sentence-ending punctuation
    last_end = max(text.rfind("."), text.rfind("!"), text.rfind("?"))
    if last_end > len(text) // 2:
        text = text[: last_end + 1]

    # Re-split into two paragraphs at the double-newline boundary
    parts = text.split("\n\n", 1)
    p1 = parts[0].strip()
    p2 = parts[1].strip() if len(parts) > 1 else ""
    return p1, p2


def _tex_to_txt(tex: str) -> str:
    """Strip LaTeX commands from .tex source to produce plain text."""
    txt = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", tex)  # \cmd{content} -> content
    txt = re.sub(r"\\[a-zA-Z]+", " ", txt)                # standalone \cmd -> space
    txt = re.sub(r"[{}]", "", txt)                         # bare braces
    txt = re.sub(r"%[^\n]*", "", txt)                      # LaTeX comments
    txt = re.sub(r"\$\|?\$", " | ", txt)                   # separators
    txt = re.sub(r"\s{2,}", " ", txt)                      # collapse whitespace
    txt = re.sub(r"\n{3,}", "\n\n", txt)                   # collapse blank lines
    return txt.strip()


def _write_docx(para1: str, para2: str, header: dict, output_path: str) -> None:
    """Write cover letter as .docx via python-docx. Silently skips if not installed."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore

        doc = Document()
        # Header block
        doc.add_heading(header.get("name", ""), level=1)
        contact_parts = filter(None, [
            header.get("email", ""),
            header.get("phone", ""),
            header.get("linkedin", ""),
        ])
        doc.add_paragraph(" | ".join(contact_parts))
        doc.add_paragraph("")
        doc.add_paragraph(para1)
        doc.add_paragraph(para2)
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
        doc.save(output_path)
    except ImportError:
        pass  # python-docx not installed; caller checks docx_path is None


def _assemble(para1: str, para2: str, header: dict, method_used: str) -> CoverLetterResult:
    """Build CoverLetterResult from two plain-text paragraphs."""
    # Build LaTeX
    template = _TEMPLATE_PATH.read_text(encoding="utf-8") if _TEMPLATE_PATH.exists() else _MINIMAL_TEX
    name = escape(header.get("name", ""))
    email = escape(header.get("email", ""))
    phone = escape(header.get("phone", ""))
    linkedin = header.get("linkedin", "")
    linkedin_text = escape(linkedin.replace("https://", "").replace("http://", ""))

    tex = template
    tex = tex.replace("{{NAME}}", name)
    tex = tex.replace("{{EMAIL}}", email)
    tex = tex.replace("{{PHONE}}", phone)
    tex = tex.replace("{{LINKEDIN}}", linkedin)
    tex = tex.replace("{{LINKEDIN_TEXT}}", linkedin_text)
    tex = tex.replace("{{PARA_ONE}}", escape(para1))
    tex = tex.replace("{{PARA_TWO}}", escape(para2))

    txt = f"{para1}\n\n{para2}"
    word_count = len(txt.split())

    # DOCX — written to a temp-adjacent path
    docx_path: Optional[str] = None
    try:
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        _write_docx(para1, para2, header, tmp.name)
        from pathlib import Path as _P
        if _P(tmp.name).stat().st_size > 0:
            docx_path = tmp.name
    except Exception:
        docx_path = None

    return CoverLetterResult(
        tex=tex,
        txt=txt,
        docx_path=docx_path,
        method_used=method_used,
        word_count=word_count,
    )


# Minimal fallback if template file is not found (e.g., during tests without full checkout)
_MINIMAL_TEX = r"""\documentclass[letterpaper,11pt]{article}
\begin{document}
\begin{center}{\Large \textbf{{{NAME}}}}\end{center}
{{EMAIL}} | {{PHONE}} | {{LINKEDIN_TEXT}}

\noindent {{PARA_ONE}}

\noindent {{PARA_TWO}}
\end{document}
"""
