"""
latex_renderer.py
Renders a LaTeX resume from a canonical profile dict + template.
All PII is runtime-injected — never hardcoded.

Template placeholders use {{KEY}} syntax.

Usage:
    python latex_renderer.py --profile profile.json --template ../templates/resume_template.tex --output resume.tex
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List


# ---------------------------------------------------------------------------
# LaTeX escaping
# ---------------------------------------------------------------------------
_LATEX_SPECIAL = {
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\^{}",
    "\\": r"\textbackslash{}",
}


def escape(text: str) -> str:
    """Escape special LaTeX characters in plain text."""
    return "".join(_LATEX_SPECIAL.get(c, c) for c in text)


def escape_url(url: str) -> str:
    """URLs go inside href — only escape % and # outside the url arg."""
    return url


# ---------------------------------------------------------------------------
# Profile → LaTeX block builders
# ---------------------------------------------------------------------------

BULLET_WORD_LIMIT = 20  # hard 2-line limit at 11pt in standard resume layout


def truncate_to_limit(text: str, limit: int = BULLET_WORD_LIMIT) -> str:
    """
    Truncate bullet text to `limit` words.
    Walks back up to 3 words to find a natural punctuation boundary (,;))
    before appending ellipsis. Preserves the full text when under the limit.
    """
    words = text.split()
    if len(words) <= limit:
        return text
    chunk = words[:limit]
    for i in range(len(chunk) - 1, max(len(chunk) - 4, 0), -1):
        if chunk[i].endswith((",", ";", ")")):
            return " ".join(chunk[:i + 1])
    return " ".join(chunk) + "..."


def render_bullets(bullets: List[Dict]) -> str:
    lines = ["    \\resumeItemListStart"]
    for b in bullets[:6]:  # max 6 bullets per role
        raw = b.get("text", "")
        text = escape(truncate_to_limit(raw))
        lines.append(f"      \\resumeItem{{{text}}}")
    lines.append("    \\resumeItemListEnd")
    return "\n".join(lines)


def render_experience(experience: List[Dict]) -> str:
    blocks = ["\\section{Experience}", "  \\resumeSubHeadingListStart", ""]
    for role in experience:
        title = escape(role.get("title", ""))
        company = escape(role.get("company", ""))
        start = escape(role.get("start", ""))
        end = escape(role.get("end", "Present"))
        location = escape(role.get("location", ""))
        date_range = f"{start} -- {end}" if start else end

        blocks.append(
            f"    \\resumeSubheading\n"
            f"      {{{title}}}{{{date_range}}}\n"
            f"      {{{company}}}{{{location}}}"
        )
        bullets = role.get("bullets", [])
        if bullets:
            blocks.append(render_bullets(bullets))
        blocks.append("")

    blocks.append("  \\resumeSubHeadingListEnd")
    return "\n".join(blocks)


def render_projects(projects: List[Dict]) -> str:
    if not projects:
        return ""
    blocks = ["\\section{Projects}", "    \\resumeSubHeadingListStart", ""]
    for proj in projects:
        name = escape(proj.get("name", ""))
        tech = escape(", ".join(proj.get("tech", [])))
        date = escape(proj.get("date", ""))
        heading = f"\\textbf{{{name}}} $|$ \\emph{{{tech}}}"
        blocks.append(
            f"      \\resumeProjectHeading\n"
            f"          {{{heading}}}{{{date}}}"
        )
        bullets = proj.get("bullets", [])
        if bullets:
            blocks.append(render_bullets(bullets))
        blocks.append("")
    blocks.append("    \\resumeSubHeadingListEnd")
    return "\n".join(blocks)


def render_skills(skills_data) -> str:
    """
    skills_data: either a list of strings or a dict with categories.
    """
    if not skills_data:
        return ""
    if isinstance(skills_data, list):
        # Plain list — output as a single skills line
        skill_str = escape(", ".join(skills_data))
        return (
            "\\section{Technical Skills}\n"
            " \\begin{itemize}[leftmargin=0.15in, label={}]\n"
            "    \\small{\\item{\n"
            f"     \\textbf{{Skills}}{{: {skill_str}}}\n"
            "    }}\n"
            " \\end{itemize}"
        )

    if isinstance(skills_data, dict):
        lines = [
            "\\section{Technical Skills}",
            " \\begin{itemize}[leftmargin=0.15in, label={}]",
            "    \\small{\\item{",
        ]
        items = []
        for category, values in skills_data.items():
            val_str = escape(", ".join(values) if isinstance(values, list) else str(values))
            items.append(f"     \\textbf{{{escape(category)}}}{{: {val_str}}}")
        lines.append(" \\\\\n".join(items))
        lines += ["    }}", " \\end{itemize}"]
        return "\n".join(lines)

    return ""


def render_education(education: List[Dict]) -> str:
    if not education:
        return ""
    blocks = ["\\section{Education}", "  \\resumeSubHeadingListStart"]
    for edu in education:
        school = escape(edu.get("school", edu.get("institution", "")))
        location = escape(edu.get("location", ""))
        degree = escape(edu.get("degree", ""))
        dates = escape(edu.get("dates", edu.get("date", "")))
        blocks.append(
            f"    \\resumeSubheading\n"
            f"      {{{school}}}{{{location}}}\n"
            f"      {{{degree}}}{{{dates}}}"
        )
    blocks.append("  \\resumeSubHeadingListEnd")
    return "\n".join(blocks)


def render_certifications(certs: List[str]) -> str:
    if not certs:
        return ""
    cert_str = " $|$ ".join(escape(c) for c in certs)
    return (
        "\\section{Certifications}\n"
        " \\begin{itemize}[leftmargin=0.15in, label={}]\n"
        f"    \\small{{\\item{{ {cert_str} }}}}\n"
        " \\end{itemize}"
    )


# ---------------------------------------------------------------------------
# Template renderer
# ---------------------------------------------------------------------------

def render_template(template_path: str, output_path: str, replacements: Dict[str, str]) -> None:
    """Replace {{KEY}} placeholders in template with rendered values."""
    content = Path(template_path).read_text(encoding="utf-8")
    for key, value in replacements.items():
        content = content.replace(f"{{{{{key}}}}}", value)

    # Clean up empty contact-line entries when LinkedIn / GitHub / portfolio
    # are not provided. Without this the header renders as "email | |" with
    # dangling separators around empty \href{}{\underline{}} fragments.
    content = re.sub(r"\\href\{\}\{\\underline\{\}\}\s*\$\|\$\s*", "", content)
    content = re.sub(r"\s*\$\|\$\s*\\href\{\}\{\\underline\{\}\}", "", content)
    content = re.sub(r"\\href\{\}\{\\underline\{\}\}", "", content)

    # Warn about any remaining unfilled placeholders
    remaining = re.findall(r"\{\{[A-Z_]+\}\}", content)
    if remaining:
        print(f"[WARNING] Unfilled placeholders: {remaining}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(content, encoding="utf-8")
    print(f"[OK] Resume written to: {output_path}")


def build_from_profile(
    profile: Dict,
    template_path: str = "../templates/resume_template.tex",
    output_path: str = "resume.tex",
    header: Dict | None = None,
) -> None:
    """
    Build a complete resume.tex from a profile dict.

    header dict keys (all runtime-provided, never hardcoded):
        name, phone, email, linkedin, github, portfolio
    """
    h = header or {}

    replacements: Dict[str, str] = {
        "NAME": escape(h.get("name", "Your Name")),
        "PHONE": escape(h.get("phone", "")),
        # EMAIL: display text (inside \underline{}) must be escaped; the href URL
        # arg is left raw so hyperref handles it correctly.
        "EMAIL": escape(h.get("email", "")),
        "LINKEDIN_URL": h.get("linkedin", ""),
        "LINKEDIN_DISPLAY": escape(h.get("linkedin", "").replace("https://", "")),
        "GITHUB_URL": h.get("github", ""),
        "GITHUB_DISPLAY": escape(h.get("github", "").replace("https://", "")),
        "PORTFOLIO_URL": h.get("portfolio", ""),
        "PORTFOLIO_DISPLAY": escape(h.get("portfolio", "").replace("https://", "")),
        "EDUCATION_SECTION": render_education(profile.get("education", [])),
        "EXPERIENCE_SECTION": render_experience(profile.get("experience", [])),
        "PROJECTS_SECTION": render_projects(profile.get("projects", [])),
        "SKILLS_SECTION": render_skills(profile.get("skills", [])),
        "CERTIFICATIONS_SECTION": render_certifications(profile.get("certifications", [])),
        "SUMMARY": escape(profile.get("summary", "")),
        # Fix 3+13: SUMMARY_SECTION renders a full \section{Summary} block when
        # a summary is present; collapses to empty string when absent so no
        # blank section appears in the final PDF.
        "SUMMARY_SECTION": (
            "\\section{Summary}\n" + escape(profile.get("summary", ""))
            if profile.get("summary", "").strip()
            else ""
        ),
    }

    render_template(template_path, output_path, replacements)


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def build_docx_from_profile(profile: Dict, output_path: str = "resume.docx", header: Dict | None = None) -> None:
    """Build a .docx resume from a canonical profile dict using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h = header or {}
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header: Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(h.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Contact line
    contact_parts = [p for p in [h.get("email", ""), h.get("phone", ""), h.get("linkedin", ""), h.get("github", "")] if p]
    if contact_parts:
        cp = doc.add_paragraph(" | ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)

    def add_section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        # Add a bottom border via XML (simple approach)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Summary
    summary = profile.get("summary", "")
    if summary and summary.strip():
        add_section_heading("Summary")
        doc.add_paragraph(summary.strip()).runs[0].font.size = Pt(9)

    # Experience
    experience = profile.get("experience", [])
    if experience:
        add_section_heading("Experience")
        for role in experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            title_run = p.add_run(role.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(10)
            company = role.get("company", "")
            start = role.get("start", "")
            end = role.get("end", "Present")
            dates = f"{start} – {end}" if start else end
            right_run = p.add_run(f"  {company}  |  {dates}")
            right_run.font.size = Pt(9)
            for bullet in role.get("bullets", [])[:6]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet.get("text", "")).font.size = Pt(9)
                bp.paragraph_format.left_indent = Inches(0.25)

    # Education
    education = profile.get("education", [])
    if education:
        add_section_heading("Education")
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            inst = edu.get("institution", edu.get("school", ""))
            deg = edu.get("degree", "")
            dates = edu.get("dates", edu.get("date", ""))
            inst_run = p.add_run(inst)
            inst_run.bold = True
            inst_run.font.size = Pt(10)
            if deg or dates:
                p.add_run(f"  |  {deg}  {dates}").font.size = Pt(9)

    # Skills
    skills = profile.get("skills", [])
    if skills:
        add_section_heading("Skills")
        skill_text = ", ".join(skills) if isinstance(skills, list) else str(skills)
        doc.add_paragraph(skill_text).runs[0].font.size = Pt(9)

    # Projects
    projects = profile.get("projects", [])
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.add_run(proj.get("name", "")).bold = True
            p.runs[-1].font.size = Pt(10)
            tech = ", ".join(proj.get("tech", []))
            if tech:
                p.add_run(f"  |  {tech}").font.size = Pt(9)
            for bullet in proj.get("bullets", [])[:4]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet.get("text", "")).font.size = Pt(9)
                bp.paragraph_format.left_indent = Inches(0.25)

    from pathlib import Path
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] DOCX resume written to: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="Render LaTeX resume from profile JSON.")
    parser.add_argument("--profile", required=True, help="Path to profile JSON")
    parser.add_argument(
        "--template",
        default=str(Path(__file__).parent.parent / "templates" / "resume_template.tex"),
        help="Path to LaTeX template",
    )
    parser.add_argument("--output", default="resume.tex", help="Output .tex path")
    parser.add_argument("--name", default="", help="Full name (runtime PII)")
    parser.add_argument("--phone", default="", help="Phone (runtime PII)")
    parser.add_argument("--email", default="", help="Email (runtime PII)")
    parser.add_argument("--linkedin", default="", help="LinkedIn URL (runtime PII)")
    parser.add_argument("--github", default="", help="GitHub URL (runtime PII)")
    parser.add_argument("--portfolio", default="", help="Portfolio URL (runtime PII)")
    parser.add_argument("--docx", action="store_true", help="Emit a .docx file instead of (or in addition to) .tex")
    args = parser.parse_args()

    with open(args.profile, encoding="utf-8") as f:
        profile = json.load(f)

    header = {
        "name": args.name,
        "phone": args.phone,
        "email": args.email,
        "linkedin": args.linkedin,
        "github": args.github,
        "portfolio": args.portfolio,
    }

    if args.docx:
        docx_output = args.output if args.output.endswith(".docx") else str(Path(args.output).with_suffix(".docx"))
        build_docx_from_profile(profile, docx_output, header)
    else:
        build_from_profile(profile, args.template, args.output, header)


if __name__ == "__main__":
    main()
